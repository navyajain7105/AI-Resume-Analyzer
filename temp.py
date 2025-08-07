import os
import uuid
import re
import fitz  # PyMuPDF
import docx
import pandas as pd
import json
import streamlit as st
from io import BytesIO
import time
from datetime import datetime

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from langchain_groq import ChatGroq
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
from dotenv import load_dotenv

# Disable LangChain tracing for better performance
os.environ["LANGCHAIN_TRACING_V2"] = "false"
os.environ["LANGCHAIN_API_KEY"] = ""

# ================= CONFIGURATION =================
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200
TOP_K = 10
MAX_RETRIES = 3

# ================= ENHANCED SCHEMAS =================
class ResumeAnalysis(BaseModel):
    domain: str = Field(..., description="Predicted domain/field of expertise")
    summary: str = Field(..., description="Comprehensive 3-4 line professional summary")
    strengths: List[str] = Field(..., description="Key strengths with specific evidence")
    weaknesses: List[str] = Field(..., description="Areas for improvement with actionable feedback")
    technical_skills: List[str] = Field(default=[], description="Verified technical skills with evidence")
    experience_level: str = Field(..., description="Junior/Mid-level/Senior based on actual experience")
    score: int = Field(..., description="Overall score out of 100 based on objective criteria")
    recommendations: List[str] = Field(default=[], description="Specific improvement recommendations")

class JobMatchAnalysis(BaseModel):
    overall_match: str = Field(..., description="yes/no/partial match assessment")
    match_percentage: int = Field(..., description="Percentage match with job requirements")
    matching_skills: List[str] = Field(default=[], description="Skills that match job requirements")
    missing_skills: List[str] = Field(default=[], description="Critical skills missing for the role")
    experience_gap: str = Field(..., description="Experience level gap analysis")
    recommendation: str = Field(..., description="Hiring recommendation with reasoning")

# ================= IMPROVED PROMPTS =================
resume_analysis_prompt = ChatPromptTemplate.from_messages([
    ("system", """You are an expert resume evaluator with 15+ years of experience in technical recruiting.

EVALUATION CRITERIA:
- Actual work experience and internships (40% weight)
- Technical projects with measurable impact (30% weight)  
- Skills backed by evidence/context (20% weight)
- Education and certifications (10% weight)

SCORING GUIDELINES:
- 90-100: Exceptional candidate with proven track record
- 80-89: Strong candidate with solid experience
- 70-79: Good candidate with relevant experience
- 60-69: Decent candidate with some gaps
- 50-59: Weak candidate, major improvements needed
- Below 50: Not suitable for most technical roles

Be fair but thorough. Focus on substance over fluff."""),
    
    ("human", """Analyze this resume comprehensively:

{resume_text}

Provide detailed analysis in this exact JSON format:
{format_instructions}

Focus on EVIDENCE-BASED evaluation. If claims lack proof, note it in weaknesses.""")
])

job_match_prompt = ChatPromptTemplate.from_messages([
    ("system", """You are a technical hiring manager evaluating resume-job fit.

MATCHING CRITERIA:
- Required technical skills coverage
- Experience level alignment  
- Domain expertise relevance
- Project complexity match
- Growth potential assessment

Be precise and objective. Provide actionable insights."""),
    
    ("human", """Compare this resume against the job requirements:

JOB DESCRIPTION:
{job_description}

RESUME:
{resume_text}

Provide detailed match analysis in JSON format:
{format_instructions}

Focus on specific skills, experience gaps, and concrete recommendations.""")
])

# ================= ENHANCED PARSERS =================
resume_parser = PydanticOutputParser(pydantic_object=ResumeAnalysis)
job_match_parser = PydanticOutputParser(pydantic_object=JobMatchAnalysis)

load_dotenv()

# Initialize LLM with better configuration
@st.cache_resource
def get_llm():
    return ChatGroq(
        model_name="llama-3.1-70b-versatile",  # More capable model
        temperature=0.1,  # Lower temperature for more consistent results
        max_retries=3
    )

llm = get_llm()

# ================= ENHANCED HELPER FUNCTIONS =================

def extract_text_robust(file) -> str:
    """Enhanced text extraction with better error handling"""
    try:
        if file.name.endswith(".pdf"):
            doc = fitz.open(stream=file.read(), filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
                # Also try OCR if text is sparse
                if len(text.strip()) < 100:
                    pix = page.get_pixmap()
                    # Could add OCR here if needed
            doc.close()
            return text
            
        elif file.name.endswith(".docx"):
            doc = docx.Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            return text
            
        elif file.name.endswith(".txt"):
            return file.read().decode("utf-8")
            
    except Exception as e:
        st.error(f"Error extracting text from {file.name}: {str(e)}")
        return ""
    
    return ""

def extract_comprehensive_metadata(text: str) -> Dict[str, Any]:
    """Extract comprehensive metadata from resume text"""
    metadata = {}
    
    # Basic contact info
    metadata["email"] = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    metadata["email"] = metadata["email"].group() if metadata["email"] else None
    
    metadata["phone"] = re.search(r"(\+91[-\s]?)?[6-9]\d{9}", text)
    metadata["phone"] = metadata["phone"].group() if metadata["phone"] else None
    
    # Social profiles
    linkedin_match = re.search(r"(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9\-_/]+", text, re.IGNORECASE)
    metadata["linkedin"] = linkedin_match.group() if linkedin_match else None
    
    github_match = re.search(r"(https?://)?(www\.)?github\.com/[a-zA-Z0-9\-_/]+", text, re.IGNORECASE)
    metadata["github"] = github_match.group() if github_match else None
    
    # Extract name (improved logic)
    name_patterns = [
        r"^([A-Z][a-zA-Z\s]+)(?:\n|$)",
        r"Name\s*[:\-]?\s*([A-Z][a-zA-Z\s]+)",
        r"([A-Z][A-Z\s]+)(?:\n.*?(?:Engineer|Developer|Analyst|Student))"
    ]
    
    for pattern in name_patterns:
        name_match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        if name_match:
            metadata["name"] = name_match.group(1).strip()
            break
    else:
        metadata["name"] = None
    
    # Extract experience level indicators
    experience_indicators = {
        "years_experience": len(re.findall(r"\b(\d+)[\+\-]?\s*years?\s*(?:of\s*)?experience", text, re.IGNORECASE)),
        "internships": len(re.findall(r"\bintern\b", text, re.IGNORECASE)),
        "projects": len(re.findall(r"\bproject\b", text, re.IGNORECASE)),
        "certifications": len(re.findall(r"\bcertificat\w*\b", text, re.IGNORECASE))
    }
    
    metadata.update(experience_indicators)
    
    return metadata

def intelligent_domain_detection(text: str, filename: str) -> str:
    """Intelligently detect resume domain based on content analysis"""
    text_lower = text.lower()
    filename_lower = filename.lower()
    
    # Domain keywords with weights
    domain_keywords = {
        "data_science": [
            "machine learning", "data science", "python", "pandas", "numpy", 
            "tensorflow", "pytorch", "sklearn", "jupyter", "data analysis",
            "statistics", "regression", "classification", "deep learning", "ai"
        ],
        "web_development": [
            "react", "angular", "vue", "javascript", "html", "css", "node.js",
            "express", "mongodb", "mysql", "rest api", "frontend", "backend",
            "full stack", "web development", "bootstrap", "jquery"
        ],
        "mobile_development": [
            "android", "ios", "react native", "flutter", "swift", "kotlin",
            "mobile app", "app development", "firebase", "xamarin"
        ],
        "devops": [
            "docker", "kubernetes", "jenkins", "ci/cd", "aws", "azure", "gcp",
            "terraform", "ansible", "devops", "cloud", "microservices", "linux"
        ],
        "cybersecurity": [
            "security", "penetration testing", "ethical hacking", "firewall",
            "cybersecurity", "vulnerability", "encryption", "infosec"
        ],
        "software_engineering": [
            "software engineer", "software developer", "programming", "algorithms",
            "data structures", "system design", "architecture", "coding"
        ]
    }
    
    # Calculate scores for each domain
    domain_scores = {}
    for domain, keywords in domain_keywords.items():
        score = sum(1 for keyword in keywords if keyword in text_lower)
        # Boost score if domain appears in filename
        if any(domain_word in filename_lower for domain_word in domain.split("_")):
            score *= 1.5
        domain_scores[domain] = score
    
    # Return the highest scoring domain, or 'general' if no clear match
    if max(domain_scores.values()) > 2:
        return max(domain_scores, key=domain_scores.get)
    else:
        return "general"

@st.cache_data(ttl=3600)  # Cache embeddings for 1 hour
def get_embedder():
    return HuggingFaceEmbeddings(
        model_name='sentence-transformers/all-MiniLM-L6-v2',
        model_kwargs={'device': 'cpu'}
    )

def build_vectorstore_optimized(documents: List[Document]) -> FAISS:
    """Build optimized vector store with better chunking"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunks = splitter.split_documents(documents)
    embedder = get_embedder()
    
    return FAISS.from_documents(chunks, embedder)

def safe_llm_call(prompt_template, **kwargs) -> Dict[str, Any]:
    """Safely call LLM with retry logic and error handling"""
    for attempt in range(MAX_RETRIES):
        try:
            input_prompt = prompt_template.format_prompt(**kwargs)
            response = llm.invoke(input_prompt.to_messages())
            
            # Clean and parse JSON response
            content = response.content.strip()
            
            # Remove markdown formatting
            if content.startswith('```json'):
                content = content[7:]
            if content.startswith('```'):
                content = content[3:]
            if content.endswith('```'):
                content = content[:-3]
            
            content = content.strip()
            
            # Find JSON boundaries
            start_idx = content.find('{')
            end_idx = content.rfind('}')
            
            if start_idx != -1 and end_idx != -1:
                json_str = content[start_idx:end_idx+1]
                return json.loads(json_str)
            
        except Exception as e:
            if attempt == MAX_RETRIES - 1:
                st.error(f"LLM call failed after {MAX_RETRIES} attempts: {str(e)}")
                return None
            time.sleep(1)  # Brief pause before retry
    
    return None

def analyze_resume_comprehensive(text: str, domain: str) -> tuple:
    """Comprehensive resume analysis with enhanced AI evaluation"""
    
    # Get basic metadata
    metadata = extract_comprehensive_metadata(text)
    
    # AI-powered analysis
    analysis_data = safe_llm_call(
        resume_analysis_prompt,
        resume_text=text,
        format_instructions=resume_parser.get_format_instructions()
    )
    
    if analysis_data:
        try:
            analysis = ResumeAnalysis(**analysis_data)
        except Exception as e:
            # Fallback analysis if parsing fails
            analysis = ResumeAnalysis(
                domain=domain,
                summary="Analysis parsing failed - manual review recommended",
                strengths=["Unable to parse detailed analysis"],
                weaknesses=["Technical parsing error occurred"],
                technical_skills=[],
                experience_level="Unknown",
                score=50,
                recommendations=["Manual review recommended due to parsing error"]
            )
    else:
        # Fallback analysis
        analysis = ResumeAnalysis(
            domain=domain,
            summary="AI analysis failed - manual review needed",
            strengths=["Analysis could not be completed"],
            weaknesses=["AI evaluation failed"],
            technical_skills=[],
            experience_level="Unknown",
            score=30,
            recommendations=["Retry analysis or manual review"]
        )
    
    # Generate unique resume ID
    resume_id = f"{domain}_{metadata['phone'] or uuid.uuid4().hex[:10]}"
    
    return analysis, metadata, resume_id

def evaluate_job_match(resume_text: str, job_description: str) -> JobMatchAnalysis:
    """Enhanced job matching with detailed analysis"""
    
    match_data = safe_llm_call(
        job_match_prompt,
        resume_text=resume_text,
        job_description=job_description,
        format_instructions=job_match_parser.get_format_instructions()
    )
    
    if match_data:
        try:
            return JobMatchAnalysis(**match_data)
        except Exception:
            pass
    
    # Fallback analysis
    return JobMatchAnalysis(
        overall_match="no",
        match_percentage=20,
        matching_skills=[],
        missing_skills=["Analysis failed - unable to determine skill gaps"],
        experience_gap="Unable to assess",
        recommendation="Manual review required due to analysis failure"
    )

def match_resumes_advanced(vectorstore: FAISS, job_desc: str, domain_filter: str = None, k: int = TOP_K) -> List[Document]:
    """Advanced resume matching with domain filtering"""
    
    # Get initial matches
    results = vectorstore.similarity_search(job_desc, k=k*2)  # Get more results for filtering
    
    if domain_filter and domain_filter != "general":
        # Filter by domain
        filtered_results = [
            doc for doc in results 
            if doc.metadata.get("domain", "").lower() == domain_filter.lower()
        ]
        return filtered_results[:k]
    
    return results[:k]

# ================= STREAMLIT UI ENHANCEMENTS =================

st.set_page_config(
    page_title="AI Resume Analyzer Pro",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
.main-header {
    text-align: center;
    color: #1f77b4;
    font-size: 2.5rem;
    margin-bottom: 2rem;
}
.metric-card {
    background: #f0f2f6;
    padding: 1rem;
    border-radius: 0.5rem;
    border-left: 4px solid #1f77b4;
    margin: 0.5rem 0;
}
.score-excellent { color: #28a745; font-weight: bold; }
.score-good { color: #17a2b8; font-weight: bold; }
.score-fair { color: #ffc107; font-weight: bold; }
.score-poor { color: #dc3545; font-weight: bold; }
</style>
""", unsafe_allow_html=True)

st.markdown('<h1 class="main-header">🎯 AI Resume Analyzer Pro</h1>', unsafe_allow_html=True)

# Sidebar configuration
with st.sidebar:
    st.header("⚙️ Configuration")
    
    analysis_mode = st.selectbox(
        "Analysis Mode",
        ["Basic Analysis", "Job Matching", "Batch Processing"],
        help="Choose the type of analysis to perform"
    )
    
    if analysis_mode == "Job Matching":
        domain_filter = st.selectbox(
            "Filter by Domain",
            ["All Domains", "Data Science", "Web Development", "Mobile Development", 
             "DevOps", "Cybersecurity", "Software Engineering"],
            help="Filter resumes by technical domain"
        )
    
    max_resumes = st.slider("Max Resumes to Process", 1, 50, 10)
    
    st.markdown("---")
    st.markdown("### 📋 Analysis Criteria")
    st.markdown("""
    - **Work Experience**: 40%
    - **Technical Projects**: 30%
    - **Skills & Evidence**: 20%
    - **Education**: 10%
    """)

# Main interface
col1, col2 = st.columns([2, 1])

with col1:
    uploaded_files = st.file_uploader(
        "📁 Upload Resume Files",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        help="Upload multiple resume files for batch analysis"
    )

with col2:
    if analysis_mode in ["Job Matching", "Batch Processing"]:
        job_desc = st.text_area(
            "📝 Job Description",
            height=200,
            help="Paste the complete job description for matching analysis"
        )

# Analysis execution
if st.button("🚀 Start Analysis", type="primary"):
    if not uploaded_files:
        st.warning("⚠️ Please upload at least one resume file.")
    elif analysis_mode == "Job Matching" and not job_desc.strip():
        st.warning("⚠️ Please provide a job description for matching analysis.")
    else:
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        all_results = []
        resume_documents = []
        
        total_files = min(len(uploaded_files), max_resumes)
        
        # Process each resume
        for idx, file in enumerate(uploaded_files[:max_resumes]):
            progress = (idx + 1) / total_files
            progress_bar.progress(progress)
            status_text.text(f"Processing {file.name}... ({idx + 1}/{total_files})")
            
            # Extract text
            text = extract_text_robust(file)
            if not text.strip():
                st.warning(f"⚠️ Could not extract text from {file.name}")
                continue
            
            # Detect domain
            domain = intelligent_domain_detection(text, file.name)
            
            try:
                # Comprehensive analysis
                analysis, metadata, resume_id = analyze_resume_comprehensive(text, domain)
                
                # Create document for vector store
                doc = Document(
                    page_content=text,
                    metadata={
                        "resume_id": resume_id,
                        "file_name": file.name,
                        "domain": domain,
                        **metadata
                    }
                )
                resume_documents.append(doc)
                
                # Prepare result data
                result = {
                    "Resume ID": resume_id,
                    "File Name": file.name,
                    "Name": metadata.get("name", "N/A"),
                    "Phone": metadata.get("phone", "N/A"),
                    "Email": metadata.get("email", "N/A"),
                    "LinkedIn": metadata.get("linkedin", "N/A"),
                    "GitHub": metadata.get("github", "N/A"),
                    "Domain": domain,
                    "Experience Level": analysis.experience_level,
                    "Summary": analysis.summary,
                    "Strengths": analysis.strengths,
                    "Weaknesses": analysis.weaknesses,
                    "Technical Skills": analysis.technical_skills,
                    "Score": analysis.score,
                    "Recommendations": analysis.recommendations
                }
                
                # Job matching analysis
                if analysis_mode == "Job Matching" and job_desc.strip():
                    job_match = evaluate_job_match(text, job_desc)
                    result.update({
                        "Job Match": job_match.overall_match,
                        "Match Percentage": job_match.match_percentage,
                        "Matching Skills": job_match.matching_skills,
                        "Missing Skills": job_match.missing_skills,
                        "Experience Gap": job_match.experience_gap,
                        "Hiring Recommendation": job_match.recommendation
                    })
                
                all_results.append(result)
                
            except Exception as e:
                st.error(f"❌ Error analyzing {file.name}: {str(e)}")
                continue
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        if all_results:
            st.success(f"✅ Successfully analyzed {len(all_results)} resumes!")
            
            # Display results
            st.markdown("## 📊 Analysis Results")
            
            # Summary statistics
            if len(all_results) > 1:
                col1, col2, col3, col4 = st.columns(4)
                
                scores = [r["Score"] for r in all_results]
                avg_score = sum(scores) / len(scores)
                
                with col1:
                    st.metric("Total Resumes", len(all_results))
                with col2:
                    st.metric("Average Score", f"{avg_score:.1f}/100")
                with col3:
                    excellent_count = sum(1 for s in scores if s >= 80)
                    st.metric("Excellent (80+)", excellent_count)
                with col4:
                    if analysis_mode == "Job Matching":
                        good_matches = sum(1 for r in all_results if r.get("Match Percentage", 0) >= 70)
                        st.metric("Good Matches (70%+)", good_matches)
            
            # Individual resume results
            for idx, result in enumerate(all_results):
                score = result["Score"]
                score_class = ("score-excellent" if score >= 80 else
                              "score-good" if score >= 65 else
                              "score-fair" if score >= 50 else "score-poor")
                
                match_info = ""
                if analysis_mode == "Job Matching" and "Match Percentage" in result:
                    match_info = f" | Match: {result['Match Percentage']}%"
                
                with st.expander(f"📄 {result['File Name']} - Score: {score}/100{match_info}", expanded=(idx < 3)):
                    col1, col2 = st.columns([1, 1])
                    
                    with col1:
                        st.markdown("### 👤 Personal Information")
                        info_data = {
                            "Name": result["Name"],
                            "Email": result["Email"],
                            "Phone": result["Phone"],
                            "Domain": result["Domain"],
                            "Experience Level": result["Experience Level"]
                        }
                        
                        for key, value in info_data.items():
                            if value and value != "N/A":
                                st.write(f"**{key}:** {value}")
                        
                        if result["LinkedIn"] != "N/A":
                            st.write(f"**LinkedIn:** [Profile]({result['LinkedIn']})")
                        if result["GitHub"] != "N/A":
                            st.write(f"**GitHub:** [Profile]({result['GitHub']})")
                    
                    with col2:
                        st.markdown("### 📊 Evaluation")
                        st.markdown(f'<p class="{score_class}">Overall Score: {score}/100</p>', 
                                  unsafe_allow_html=True)
                        
                        if analysis_mode == "Job Matching" and "Match Percentage" in result:
                            st.write(f"**Job Match:** {result['Job Match'].title()}")
                            st.write(f"**Match Percentage:** {result['Match Percentage']}%")
                    
                    # Summary
                    st.markdown("### 📝 Summary")
                    st.write(result["Summary"])
                    
                    # Strengths and Weaknesses
                    col3, col4 = st.columns(2)
                    
                    with col3:
                        st.markdown("### ✅ Strengths")
                        for strength in result["Strengths"]:
                            st.write(f"• {strength}")
                    
                    with col4:
                        st.markdown("### ⚠️ Areas for Improvement")
                        for weakness in result["Weaknesses"]:
                            st.write(f"• {weakness}")
                    
                    # Technical skills
                    if result["Technical Skills"]:
                        st.markdown("### 🛠️ Technical Skills")
                        skills_text = " | ".join(result["Technical Skills"])
                        st.write(skills_text)
                    
                    # Job matching details
                    if analysis_mode == "Job Matching" and "Matching Skills" in result:
                        col5, col6 = st.columns(2)
                        
                        with col5:
                            st.markdown("### ✅ Matching Skills")
                            for skill in result["Matching Skills"]:
                                st.write(f"• {skill}")
                        
                        with col6:
                            st.markdown("### ❌ Missing Skills")
                            for skill in result["Missing Skills"]:
                                st.write(f"• {skill}")
                        
                        st.markdown("### 💡 Hiring Recommendation")
                        st.write(result["Hiring Recommendation"])
                    
                    # Recommendations
                    if result["Recommendations"]:
                        st.markdown("### 🎯 Improvement Recommendations")
                        for rec in result["Recommendations"]:
                            st.write(f"• {rec}")
            
            # Export functionality
            st.markdown("---")
            st.markdown("## 📥 Export Results")
            
            # Prepare DataFrame for export
            export_data = []
            for result in all_results:
                export_row = {
                    "Resume_ID": result["Resume ID"],
                    "File_Name": result["File Name"],
                    "Name": result["Name"],
                    "Email": result["Email"],
                    "Phone": result["Phone"],
                    "LinkedIn": result["LinkedIn"],
                    "GitHub": result["GitHub"],
                    "Domain": result["Domain"],
                    "Experience_Level": result["Experience Level"],
                    "Score": result["Score"],
                    "Summary": result["Summary"],
                    "Strengths": " | ".join(result["Strengths"]),
                    "Weaknesses": " | ".join(result["Weaknesses"]),
                    "Technical_Skills": " | ".join(result["Technical Skills"]),
                    "Recommendations": " | ".join(result["Recommendations"])
                }
                
                if analysis_mode == "Job Matching" and "Match Percentage" in result:
                    export_row.update({
                        "Job_Match": result["Job Match"],
                        "Match_Percentage": result["Match Percentage"],
                        "Matching_Skills": " | ".join(result["Matching Skills"]),
                        "Missing_Skills": " | ".join(result["Missing Skills"]),
                        "Experience_Gap": result["Experience Gap"],
                        "Hiring_Recommendation": result["Hiring Recommendation"]
                    })
                
                export_data.append(export_row)
            
            df = pd.DataFrame(export_data)
            
            # Create Excel file
            excel_buffer = BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Resume_Analysis', index=False)
                
                # Add summary sheet if multiple resumes
                if len(all_results) > 1:
                    summary_data = {
                        "Metric": ["Total Resumes", "Average Score", "Excellent (80+)", "Good (65-79)", "Fair (50-64)", "Poor (<50)"],
                        "Value": [
                            len(all_results),
                            f"{avg_score:.1f}",
                            sum(1 for s in scores if s >= 80),
                            sum(1 for s in scores if 65 <= s < 80),
                            sum(1 for s in scores if 50 <= s < 65),
                            sum(1 for s in scores if s < 50)
                        ]
                    }
                    
                    summary_df = pd.DataFrame(summary_data)
                    summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            excel_buffer.seek(0)
            
            st.download_button(
                label="📊 Download Detailed Excel Report",
                data=excel_buffer,
                file_name=f"resume_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            
        else:
            st.error("❌ No resumes were successfully processed. Please check your files and try again.")

# Footer
st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: #666;'>"
    "Built with ❤️ using Streamlit & LangChain | "
    "AI-Powered Resume Analysis"
    "</div>",
    unsafe_allow_html=True
)
