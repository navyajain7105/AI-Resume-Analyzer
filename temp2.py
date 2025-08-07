import os
import uuid
import re
import fitz  # PyMuPDF
import docx
import pandas as pd
import json
import streamlit as st
from io import BytesIO

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from langchain_groq import ChatGroq
from pydantic import BaseModel, Field
from typing import List, Optional
from dotenv import load_dotenv

os.environ["LANGCHAIN_TRACING_V2"] = "false"
# Make sure to set your Groq API key in your environment variables or replace ""
os.environ["GROQ_API_KEY"] = "" 

# ------------------- Config -------------------
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 100
TOP_K = 5

# ------------------- Schema Definition -------------------
class ResumeAnalysis(BaseModel):
    domain: str = Field(..., description="Predicted domain of the resume")
    summary: Optional[str] = Field(..., description="3-line summary")
    strengths: List[str] = Field(..., description="Resume strengths as a list")
    weaknesses: List[str] = Field(..., description="Resume weaknesses as a list")
    score: int = Field(..., description="Overall score out of 100")
    job_match: Optional[str] = Field(default="no", description="Whether resume matches job requirements")

parser = PydanticOutputParser(pydantic_object=ResumeAnalysis)

resume_prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a professional resume evaluator."),
    ("human", """Analyze the following resume and return this JSON:

{format_instructions}

Resume:
{resume_text}
""")
])

# IMPROVED EVALUATION PROMPT - More realistic and domain-specific
evaluation_prompt = ChatPromptTemplate.from_messages([
    ("system", 
     "You are a strict but fair technical recruiter. You must respond with ONLY valid JSON, no other text.\n\n"
     "REALISTIC SCORING APPROACH:\n"
     "- Entry-level candidates (students) should typically score 40-70, not 80+\n"
     "- Domain relevance is CRITICAL - irrelevant projects should not boost scores significantly\n"
     "- Academic projects are valuable but less than real-world experience\n"
     "- High CGPA shows learning ability but doesn't replace domain expertise\n\n"
     "DOMAIN-SPECIFIC EVALUATION:\n"
     "- For ML/Data Science: Look for ML algorithms, data analysis, statistics, model deployment\n"
     "- For Software Development: Look for web frameworks, databases, system design, deployment\n"
     "- For DevOps: Look for CI/CD, containerization, cloud platforms, infrastructure\n"
     "- Cross-domain projects have limited relevance (max 20% weight)\n\n"
     "SCORING SCALE:\n"
     "- 80-100: Exceptional candidates with strong domain expertise and experience\n"
     "- 65-79: Good candidates with solid domain skills and some experience\n"
     "- 50-64: Average candidates with basic domain knowledge\n"
     "- 35-49: Below average with limited domain relevance\n"
     "- Below 35: Significant skills gap for the specific domain\n\n"
     "EXPERIENCE WEIGHT:\n"
     "- Real industry experience: High weight\n"
     "- Internships in relevant domain: Medium-high weight\n"
     "- Academic projects in domain: Medium weight\n"
     "- Personal projects in domain: Medium weight\n"
     "- Cross-domain projects: Low weight\n"
     "- Teaching/mentoring: Shows communication but limited technical weight\n\n"
     "BE REALISTIC: Don't inflate scores. Better to be conservative and accurate."
    ),
     ("human", 
       """Analyze this resume against the job description with STRICT domain relevance.

Job Description:
{jd_text}

Resume:
{resume_text}

CRITICAL EVALUATION CRITERIA:
1. Domain Relevance (40% weight): How well do skills/projects match the specific domain?
2. Technical Depth (30% weight): Complexity and depth of technical implementation
3. Real-world Experience (20% weight): Industry internships, real deployments, production code
4. Learning Potential (10% weight): Academic performance, growth trajectory

DOMAIN-SPECIFIC REQUIREMENTS:
- ML/Data Science: Machine learning algorithms, data preprocessing, model evaluation, deployment
- Software Development: Full-stack development, databases, APIs, system architecture
- DevOps: CI/CD pipelines, containerization, cloud infrastructure, monitoring
- Backend: Server-side development, databases, APIs, scalability
- Frontend: UI/UX, responsive design, modern frameworks, user experience

SCORING RULES:
- If candidate lacks 70%+ domain-relevant experience: Score ‚â§ 55
- If candidate has only academic projects: Score ‚â§ 65
- If candidate has cross-domain skills only: Score ‚â§ 45
- Reserve 80+ scores for candidates with strong domain expertise + real experience

Return ONLY this JSON format:
{{
  "domain": "predicted domain",
  "summary": "honest assessment focusing on domain fit",
  "strengths": ["domain-relevant strengths with evidence"],
  "weaknesses": ["honest areas lacking for this specific role"],
  "score": 45
}}

JSON response:""")
])

load_dotenv()

# Ensure you have set the GROQ_API_KEY environment variable
llm = ChatGroq(model_name="llama-3.1-8b-instant")

# ------------------------- Helper Functions -------------------------

def extract_text(file):
    if file.name.endswith(".pdf"):
        try:
            # Clean up text by joining lines that might have been split
            doc = fitz.open(stream=file.read(), filetype="pdf")
            full_text = ""
            for page in doc:
                full_text += page.get_text()
            # A simple way to join hyphenated words split across lines
            full_text = re.sub(r"-\n", "", full_text)
            return full_text
        except:
            return ""
    elif file.name.endswith(".docx"):
        try:
            doc = docx.Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        except:
            return ""
    elif file.name.endswith(".txt"):
        try:
            return file.read().decode("utf-8")
        except:
            return ""
    return ""

def extract_metadata(text: str):
    """Improved metadata extraction with better regex patterns"""
    # Join lines to handle multi-line contact info before regex matching
    cleaned_text = re.sub(r'\s*\n\s*', ' ', text)

    # Name extraction - look for patterns after common indicators
    name_patterns = [
        r"(?i)^([A-Z][a-z]+ [A-Z][a-z]+)",  # First line capitalized names
        r"(?i)name\s*[:\-]?\s*([A-Za-z\s]+)",  # "Name:" pattern
        r"(?i)^([A-Z\s]{2,30})\s*$",  # All caps names on single line
    ]
    
    name = None
    # Use original text for name finding as it relies on line structure
    for pattern in name_patterns:
        name_match = re.search(pattern, text, re.MULTILINE)
        if name_match:
            candidate_name = name_match.group(1).strip()
            # Filter out common false positives
            if candidate_name and len(candidate_name) > 3 and not any(word in candidate_name.lower() for word in ['contact', 'education', 'experience', 'skills']):
                name = candidate_name
                break
    
    # Phone extraction - support various formats
    phone_patterns = [
        r"\b(\+?91[-\s]?)?([6789]\d{9})\b",  # Indian format
        r"\b(\d{10})\b",  # 10 digit
        r"\+\d{1,3}[-.\s]?\d{10,}",  # International format
    ]
    
    phone = None
    for pattern in phone_patterns:
        phone_match = re.search(pattern, cleaned_text)
        if phone_match:
            phone = phone_match.group().strip()
            break
    
    # Email extraction
    email_match = re.search(r"[\w\.\-]+@[\w\.\-]+\.\w+", cleaned_text)
    
    # LinkedIn extraction - improved pattern
    linkedin_patterns = [
        r"(https?://)?(www\.)?linkedin\.com/in/[\w\-\/]+",
    ]
    
    linkedin = None
    for pattern in linkedin_patterns:
        linkedin_match = re.search(pattern, cleaned_text, re.IGNORECASE)
        if linkedin_match:
            linkedin_url = linkedin_match.group().strip()
            # Ensure it's a proper URL
            if not linkedin_url.startswith('http'):
                linkedin_url = 'https://' + linkedin_url
            linkedin = linkedin_url
            break
    
    # GitHub extraction - improved pattern
    github_patterns = [
        r"(https?://)?(www\.)?github\.com/[\w\-\/]+",
    ]
    
    github = None
    for pattern in github_patterns:
        github_match = re.search(pattern, cleaned_text, re.IGNORECASE)
        if github_match:
            github_url = github_match.group().strip()
            # Ensure it's a proper URL
            if not github_url.startswith('http'):
                github_url = 'https://' + github_url
            github = github_url
            break

    return {
        "name": name,
        "phone": phone,
        "email": email_match.group() if email_match else None,
        "linkedin": linkedin,
        "github": github,
    }

def get_embedder():
    return HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')

def build_vectorstore(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = splitter.split_documents(documents)
    embedder = get_embedder()
    return FAISS.from_documents(chunks, embedder)

def match_resumes(vectorstore, job_desc, domain_filter, k=TOP_K):
    results = vectorstore.similarity_search(job_desc, k=15)
    filtered = [doc for doc in results if doc.metadata.get("domain") == domain_filter.lower()]
    return filtered[:k]

def analyze_resume(text, domain):
    input_prompt = resume_prompt.format_prompt(
        resume_text=text,
        format_instructions=parser.get_format_instructions()
    )
    output = llm.invoke(input_prompt.to_messages())
    analysis = parser.parse(output.content)
    meta = extract_metadata(text)
    resume_id = f"{domain}_{meta['phone'] or uuid.uuid4().hex[:10]}"
    return analysis, meta, resume_id

def clean_json_response(raw_response: str) -> str:
    """Clean and extract JSON from LLM response"""
    raw_response = raw_response.strip()
    
    if raw_response.startswith('```json'):
        raw_response = raw_response[7:]
    if raw_response.startswith('```'):
        raw_response = raw_response[3:]
    if raw_response.endswith('```'):
        raw_response = raw_response[:-3]
    
    raw_response = raw_response.strip()
    
    start = raw_response.find('{')
    if start == -1:
        raise ValueError(f"No opening brace found in response: '{raw_response[:100]}...'")
    
    brace_count = 0
    end = -1
    for i in range(start, len(raw_response)):
        if raw_response[i] == '{':
            brace_count += 1
        elif raw_response[i] == '}':
            brace_count -= 1
            if brace_count == 0:
                end = i
                break
    
    if end == -1:
        raise ValueError(f"No matching closing brace found. Response starts with: '{raw_response[:200]}...'")
    
    json_str = raw_response[start:end+1]
    return json_str

def calculate_domain_specific_score(resume_text: str, target_domain: str) -> int:
    """
    Calculate domain-specific score with emphasis on proven work rather than claimed skills
    """
    resume_lower = resume_text.lower()
    
    # PROVEN WORK INDICATORS - These carry the most weight
    proven_work_indicators = {
        'projects': {
            'strong_evidence': ['[github.com/](https://github.com/)', 'deployed', 'live demo', 'production', 'website:', 'app store', 'play store'],
            'medium_evidence': ['project', 'developed', 'built', 'created', 'implemented', 'designed'],
            'weak_evidence': ['familiar with', 'knowledge of', 'learned']
        },
        'experience': {
            'strong_evidence': ['internship', 'work experience', 'employed', 'professional', 'industry', 'company'],
            'medium_evidence': ['freelance', 'contract', 'part-time', 'volunteer work'],
            'weak_evidence': ['teaching assistant', 'mentor', 'tutor']
        },
        'achievements': {
            'strong_evidence': ['winner', 'first place', 'awarded', 'competition winner', 'hackathon winner'],
            'medium_evidence': ['participated', 'competition', 'hackathon', 'contest'],
            'weak_evidence': ['attended', 'workshop', 'seminar']
        }
    }
    
    # Calculate proven work score (70% of total)
    proven_score = 0
    
    # Strong evidence of work (40% weight)
    strong_work_count = sum(1 for category in proven_work_indicators.values() 
                            for indicator in category['strong_evidence'] 
                            if indicator in resume_lower)
    proven_score += min(40, strong_work_count * 8)
    
    # Medium evidence of work (20% weight)
    medium_work_count = sum(1 for category in proven_work_indicators.values() 
                            for indicator in category['medium_evidence'] 
                            if indicator in resume_lower)
    proven_score += min(20, medium_work_count * 3)
    
    # Weak evidence (10% weight)
    weak_work_count = sum(1 for category in proven_work_indicators.values() 
                          for indicator in category['weak_evidence'] 
                          if indicator in resume_lower)
    proven_score += min(10, weak_work_count * 2)
    
    # DOMAIN-SPECIFIC IMPLEMENTATION EVIDENCE (20% of total)
    domain_implementation_score = 0
    
    # Define what counts as implementation evidence for each domain
    implementation_evidence = {
        'machine_learning': [
            'model accuracy', 'dataset', 'training', 'prediction', 'classification accuracy',
            'implemented neural network', 'built ml model', 'data preprocessing',
            'model deployment', 'tensorflow implementation', 'pytorch model'
        ],
        'software_development': [
            'rest api', 'database connection', 'user authentication', 'responsive design',
            'full stack application', 'web application', 'mobile app', 'backend server',
            'frontend interface', 'database schema', 'deployed application'
        ],
        'devops': [
            'ci/cd pipeline', 'automated deployment', 'infrastructure setup', 'monitoring system',
            'containerized application', 'cloud deployment', 'server configuration',
            'automated testing', 'deployment script', 'infrastructure code'
        ],
        'data_science': [
            'data visualization', 'statistical analysis', 'data cleaning', 'insights generated',
            'dashboard created', 'report generated', 'trend analysis', 'data pipeline',
            'business intelligence', 'predictive analytics'
        ]
    }
    
    # Map target domain
    domain_map = {
        'ml': 'machine_learning',
        'machine learning': 'machine_learning',
        'data science': 'data_science',
        'datascience': 'data_science',
        'software': 'software_development',
        'fullstack': 'software_development',
        'backend': 'software_development',
        'frontend': 'software_development',
        'devops': 'devops'
    }
    
    mapped_domain = domain_map.get(target_domain.lower(), 'software_development')
    relevant_implementations = implementation_evidence.get(mapped_domain, implementation_evidence['software_development'])
    
    implementation_count = sum(1 for impl in relevant_implementations if impl in resume_lower)
    domain_implementation_score = min(20, implementation_count * 4)
    
    # SKILLS AND COURSES - Heavily reduced weight (10% of total)
    skills_score = 0
    
    # Only count skills if they appear in context of actual work/projects
    skill_context_patterns = [
        r'used\s+(\w+)', r'implemented\s+(\w+)', r'built\s+with\s+(\w+)',
        r'developed\s+using\s+(\w+)', r'created\s+with\s+(\w+)'
    ]
    
    domain_skills = {
        'machine_learning': ['python', 'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy'],
        'software_development': ['javascript', 'react', 'node.js', 'python', 'java', 'mysql', 'mongodb'],
        'devops': ['docker', 'kubernetes', 'aws', 'jenkins', 'ansible', 'terraform'],
        'data_science': ['python', 'r', 'sql', 'tableau', 'pandas', 'matplotlib']
    }
    
    relevant_skills = domain_skills.get(mapped_domain, domain_skills['software_development'])
    
    # Count skills only when mentioned in context of actual work
    contextual_skills = 0
    for pattern in skill_context_patterns:
        matches = re.findall(pattern, resume_lower)
        contextual_skills += sum(1 for match in matches if match in relevant_skills)
    
    skills_score = min(10, contextual_skills * 2)
    
    # Calculate final score
    final_score = int(proven_score + domain_implementation_score + skills_score)
    
    # Apply penalties for pure skill/course listing without evidence
    skill_listing_patterns = ['skills:', 'technologies:', 'programming languages:', 'tools:']
    course_patterns = ['courses:', 'coursework:', 'subjects:', 'curriculum:']
    
    has_skill_lists = any(pattern in resume_lower for pattern in skill_listing_patterns)
    has_course_lists = any(pattern in resume_lower for pattern in course_patterns)
    
    # Penalty for resumes that are mostly skill lists without proven work
    if has_skill_lists and strong_work_count < 2:
        final_score = int(final_score * 0.8)  # 20% penalty
    
    if has_course_lists and implementation_count < 2:
        final_score = int(final_score * 0.9)  # 10% penalty
    
    # Apply realistic caps
    if 'internship' not in resume_lower and 'work experience' not in resume_lower and strong_work_count < 3:
        final_score = min(final_score, 65)  # Cap for students without strong proof of work
    
    # Minimum score for candidates with some proven work
    if strong_work_count > 0 or implementation_count > 1:
        final_score = max(final_score, 30)
    
    return final_score

def evaluate_resume_vs_jd(resume_text: str, jd_text: str):
    """Evaluate resume against job description with realistic scoring"""
    
    # Extract target domain from job description
    jd_lower = jd_text.lower()
    if any(term in jd_lower for term in ['machine learning', 'ml engineer', 'data scientist', 'ai']):
        target_domain = 'machine_learning'
    elif any(term in jd_lower for term in ['software developer', 'full stack', 'backend', 'frontend', 'web developer']):
        target_domain = 'software_development'
    elif any(term in jd_lower for term in ['devops', 'site reliability', 'infrastructure', 'deployment']):
        target_domain = 'devops'
    else:
        target_domain = 'software_development'  # default
    
    for attempt in range(3):
        try:
            input_prompt = evaluation_prompt.format_prompt(
                resume_text=resume_text,
                jd_text=jd_text
            )
            
            response = llm.invoke(input_prompt.to_messages())
            raw = response.content.strip()
            
            if len(raw) < 20 or not ('{' in raw and '}' in raw):
                if attempt < 2:
                    continue
                else:
                    raise ValueError(f"All attempts failed. Last response: '{raw}'")
            
            try:
                cleaned_json = clean_json_response(raw)
            except Exception as clean_error:
                if attempt < 2:
                    continue
                else:
                    raise ValueError(f"JSON cleaning failed after all attempts: {clean_error}")
            
            try:
                data = json.loads(cleaned_json)
                
                required_keys = ["domain", "summary", "strengths", "weaknesses", "score"]
                missing_keys = [key for key in required_keys if key not in data]
                if missing_keys:
                    if attempt < 2:
                        continue
                    else:
                        raise ValueError(f'Missing keys: {missing_keys} in parsed JSON')
                
                # Ensure correct data types
                if not isinstance(data["strengths"], list):
                    if isinstance(data["strengths"], str):
                        data["strengths"] = [s.strip() for s in data["strengths"].split('\n') if s.strip()]
                    else:
                        data["strengths"] = ["Shows learning potential and technical curiosity"]
                
                if not isinstance(data["weaknesses"], list):
                    if isinstance(data["weaknesses"], str):
                        data["weaknesses"] = [w.strip() for w in data["weaknesses"].split('\n') if w.strip()]
                    else:
                        data["weaknesses"] = ["Limited domain-specific experience"]
                
                # Calculate realistic score based on domain relevance
                try:
                    llm_score = int(float(data["score"]))
                    domain_specific_score = calculate_domain_specific_score(resume_text, target_domain)
                    
                    # Take the more conservative score
                    final_score = min(llm_score, domain_specific_score)
                    
                    # Additional penalty for domain mismatch
                    if target_domain == 'machine_learning' and 'machine learning' not in resume_text.lower():
                        final_score = min(final_score, 45)  # Cap at 45 for non-ML candidates applying to ML roles
                    elif target_domain == 'software_development' and not any(term in resume_text.lower() for term in ['web', 'software', 'application', 'system', 'backend', 'frontend']):
                        final_score = min(final_score, 40)  # Cap for non-software candidates
                    
                    data["score"] = final_score
                except (ValueError, TypeError):
                    data["score"] = calculate_domain_specific_score(resume_text, target_domain)
                
                # Set job match based on realistic threshold
                data["job_match"] = "yes" if data["score"] >= 55 else "no"
                
                result = ResumeAnalysis(**data)
                return result
                
            except json.JSONDecodeError as json_error:
                if attempt < 2:
                    continue
                else:
                    raise ValueError(f"JSON parsing failed after all attempts: {json_error}")
            
        except Exception as e:
            if attempt < 2:
                continue
            else:
                break
    
    # Fallback evaluation with domain-specific scoring
    realistic_score = calculate_domain_specific_score(resume_text, target_domain)
    
    return ResumeAnalysis(
        domain=target_domain.replace('_', ' '),
        summary=f"Entry-level candidate with foundational skills in computer science. Shows academic competence but limited domain-specific experience for {target_domain.replace('_', ' ')} roles.",
        strengths=[
            "Strong academic background with good CGPA",
            "Basic programming skills and project experience", 
            "Active learning mindset and technical curiosity",
            "Involvement in extracurricular activities"
        ],
        weaknesses=[
            f"Limited {target_domain.replace('_', ' ')} specific experience",
            "Lack of industry internships or professional experience",
            "Projects not directly aligned with target domain",
            "Need to develop deeper technical expertise"
        ],
        score=realistic_score,
        job_match="yes" if realistic_score >= 55 else "no"
    )

# ------------------- Streamlit App -------------------
st.set_page_config(page_title="AI Resume Analyzer - Improved", layout="wide")
st.title("üìÑ AI Resume Analyzer & Matcher")

st.info("üîß **Improvements:** This version uses a dual-scoring system for more realistic evaluations, has better data extraction, and provides domain-specific analysis.", icon="üîß")

uploaded_files = st.file_uploader("Upload Resumes (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
job_desc = st.text_area("Paste Job Description (Optional - for job matching)")

if st.button("üöÄ Run Analysis", use_container_width=True):
    if not uploaded_files:
        st.warning("Please upload at least one resume.")
    else:
        progress_bar = st.progress(0, text="Starting analysis...")
        all_results = []
        resume_docs = []

        for i, file in enumerate(uploaded_files):
            progress_bar.progress((i + 1) / len(uploaded_files), text=f"Analyzing: {file.name}")
            text = extract_text(file)
            if not text.strip():
                st.warning(f"Could not extract text from: {file.name}")
                continue

            # Use JD to infer domain if available, otherwise guess from filename
            if job_desc:
                jd_lower = job_desc.lower()
                if any(term in jd_lower for term in ['machine learning', 'ml', 'ai']):
                    inferred_domain = 'machine_learning'
                elif any(term in jd_lower for term in ['software', 'full stack', 'backend']):
                    inferred_domain = 'software_development'
                else:
                    inferred_domain = 'general'
            else:
                 inferred_domain = "devops" if "devops" in file.name.lower() else (
                                 "datascience" if "ds" in file.name.lower() else (
                                 "java" if "java" in file.name.lower() else (
                                 "fullstack" if "full" in file.name.lower() else "general")))

            try:
                # If JD is present, use the more advanced evaluation
                if job_desc:
                    analysis = evaluate_resume_vs_jd(text, job_desc)
                    meta = extract_metadata(text)
                    resume_id = f"{analysis.domain}_{meta['phone'] or uuid.uuid4().hex[:10]}"
                else:
                    analysis, meta, resume_id = analyze_resume(text, inferred_domain)

                doc = Document(
                    page_content=text,
                    metadata={
                        "resume_id": resume_id,
                        "file_name": file.name,
                        "domain": analysis.domain,
                        **meta
                    }
                )
                resume_docs.append(doc)

                all_results.append({
                    "Resume ID": resume_id,
                    "File Name": file.name,
                    "Name": meta["name"],
                    "Phone": meta["phone"],
                    "Email": meta["email"],
                    "LinkedIn": meta["linkedin"],
                    "GitHub": meta["github"],
                    "Domain": analysis.domain,
                    "Summary": analysis.summary,
                    "Strengths": "\n".join(f"‚Ä¢ {s}" for s in analysis.strengths),
                    "Weaknesses": "\n".join(f"‚Ä¢ {w}" for w in analysis.weaknesses),
                    "Score": analysis.score,
                    "Job Match": "‚úÖ Yes" if analysis.job_match == "yes" else "‚ùå No"
                })

            except Exception as e:
                st.error(f"Error analyzing {file.name}: {e}")
        
        progress_bar.progress(1.0, text="Analysis complete!")

        if all_results:
            df = pd.DataFrame(all_results)
            st.success("‚úÖ Analysis Complete")

            # Display each resume as an expandable section
            for idx, row in df.sort_values("Score", ascending=False).iterrows():
                if row['Score'] >= 65:
                    score_color = "green"
                    score_icon = "üü¢"
                elif row['Score'] >= 50:
                    score_color = "orange"
                    score_icon = "üü°"
                else:
                    score_color = "red"
                    score_icon = "üî¥"
                
                with st.expander(f"**{row['Name'] or row['File Name']}** | Score: **:{score_color}[{row['Score']}]** | Match: {row.get('Job Match', 'N/A')}"):
                    st.subheader("üìù Evaluation Summary")
                    st.markdown(f"**Domain:** `{row['Domain']}`")
                    st.markdown(f"**Overall Assessment:** {row['Summary']}")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("**‚úÖ Strengths**")
                        st.markdown(row['Strengths'])
                    with col2:
                        st.markdown("**‚ö†Ô∏è Weaknesses**")
                        st.markdown(row['Weaknesses'])
                    
                    st.divider()
                    st.subheader("üë§ Contact Information")
                    st.markdown(f"""
                    - **Email:** {row['Email'] or 'Not Found'}
                    - **Phone:** {row['Phone'] or 'Not Found'}
                    - **LinkedIn:** [{row['LinkedIn']}]({row['LinkedIn']}) if row['LinkedIn'] else 'Not Found'
                    - **GitHub:** [{row['GitHub']}]({row['GitHub']}) if row['GitHub'] else 'Not Found'
                    """)

            st.write("---")
            
            # Show a compact summary table
            st.write("### üìã Summary Table")
            summary_df = df[['Name', 'Score', 'Job Match', 'Domain', 'File Name']].copy()
            st.dataframe(
                summary_df.sort_values("Score", ascending=False),
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Score": st.column_config.ProgressColumn(
                        "Score",
                        format="%d",
                        min_value=0,
                        max_value=100,
                    ),
                }
            )

            @st.cache_data
            def convert_df_to_excel(df_to_convert):
                output = BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df_to_convert.to_excel(writer, index=False, sheet_name='Resume Analysis')
                processed_data = output.getvalue()
                return processed_data

            excel_data = convert_df_to_excel(df)
            st.download_button(
                label="üì• Download Full Analysis as Excel",
                data=excel_data,
                file_name="resume_analysis_results.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
