import os
import uuid
import re
import fitz  # PyMuPDF
import docx
import pandas as pd
import streamlit as st
from io import BytesIO

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from langchain_groq import ChatGroq
from pydantic import BaseModel, Field
from typing import List
from dotenv import load_dotenv

os.environ["LANGCHAIN_TRACING_V2"] = "false"
os.environ["LANGCHAIN_API_KEY"] = ""

# ------------------- Config -------------------
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 100
TOP_K = 5

# ------------------- Schema Definition -------------------
class ResumeAnalysis(BaseModel):
    domain: str = Field(..., description="Predicted domain of the resume")
    summary: str = Field(..., description="3-line summary")
    strengths: List[str] = Field(..., description="Resume strengths as a list")
    weaknesses: List[str] = Field(..., description="Resume weaknesses as a list")
    score: int = Field(..., description="Overall score out of 100")

parser = PydanticOutputParser(pydantic_object=ResumeAnalysis)

resume_prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a professional resume evaluator."),
    ("human", """Analyze the following resume and return this JSON:

{format_instructions}

Resume:
{resume_text}
""")
])

evaluation_prompt = ChatPromptTemplate.from_messages([
    ("system", 
     "You are a strict technical evaluator comparing a candidate resume against a job description.\n"
     "Focus only on *actual experience*. Do not trust skills section if no real work/project context is found.\n"
     "Check alignment of cloud platform (e.g., AWS vs Azure), tools, domain expertise, project relevance, experience years, and gaps.\n"
     "Be strict and conservative in scoring. Only proven skills matter.\n\n"
     "Return JSON with: domain, strengths, weaknesses, and score out of 100.\n\n"
     "Scoring Criteria:\n"
     "- Required experience missing = big penalty\n"
     "- Skill mentioned but no work done = medium penalty\n"
     "- Good alignment = boost\n"
     "- Irrelevant experience = neutral or penalty\n"),
    ("human", 
     "Job Description:\n{jd_text}\n\n"
     "Candidate Resume:\n{resume_text}\n\n"
     "Return output in JSON with the following keys:\n"
     "domain, strengths (list), weaknesses (list), score (integer out of 100)")
])

load_dotenv()

llm = ChatGroq(model_name="llama-3.1-8b-instant")

# ------------------------- Helper Functions -------------------------

def extract_text(file):
    if file.name.endswith(".pdf"):
        try:
            doc = fitz.open(stream=file.read(), filetype="pdf")
            return "\n".join(page.get_text() for page in doc)
        except:
            return ""
    elif file.name.endswith(".docx"):
        try:
            doc = docx.Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        except:
            return ""
    elif file.name.endswith(".txt"):
        try:
            return file.read().decode("utf-8")
        except:
            return ""
    return ""

def extract_metadata(text: str):
    name_match = re.search(r"Name\s*[:\-]?\s*(.+)", text, re.IGNORECASE)
    phone_match = re.search(r"\b\d{10}\b", text)
    email_match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    linkedin_match = re.search(r"(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9\-_/]+", text)
    github_match = re.search(r"(https?://)?(www\.)?github\.com/[a-zA-Z0-9\-_/]+", text)

    return {
        "name": name_match.group(1).strip() if name_match else None,
        "phone": phone_match.group() if phone_match else None,
        "email": email_match.group() if email_match else None,
        "linkedin": linkedin_match.group() if linkedin_match else None,
        "github": github_match.group() if github_match else None,
    }

def get_embedder():
    return HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')

def build_vectorstore(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = splitter.split_documents(documents)
    embedder = get_embedder()
    return FAISS.from_documents(chunks, embedder)

def match_resumes(vectorstore, job_desc, domain_filter, k=TOP_K):
    results = vectorstore.similarity_search(job_desc, k=15)
    filtered = [doc for doc in results if doc.metadata.get("domain") == domain_filter.lower()]
    return filtered[:k]

def analyze_resume(text, domain):
    input_prompt = resume_prompt.format_prompt(
        resume_text=text,
        format_instructions=parser.get_format_instructions()
    )
    output = llm.invoke(input_prompt.to_messages())
    analysis = parser.parse(output.content)
    meta = extract_metadata(text)
    resume_id = f"{domain}_{meta['phone'] or uuid.uuid4().hex[:10]}"
    return analysis, meta, resume_id

def evaluate_resume_vs_jd(resume_text: str, jd_text: str):
    input_prompt = evaluation_prompt.format_prompt(resume_text=resume_text, jd_text=jd_text)
    output = llm.invoke(input_prompt.to_messages())
    return parser.parse(output.content)

# ------------------- Streamlit App -------------------
st.set_page_config(page_title="AI Resume Analyzer")
st.title("Resume Analyzer + Matcher + Excel Export")

uploaded_files = st.file_uploader("Upload Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True)
job_desc = st.text_area("Paste Job Description (optional)")

if st.button("Run Analysis"):
    if not uploaded_files:
        st.warning("Please upload resumes.")
    else:
        st.info("Analyzing resumes...")
        all_results = []
        resume_docs = []

        for file in uploaded_files:
            text = extract_text(file)
            if not text.strip():
                st.warning(f"Could not extract text from: {file.name}")
                continue

            inferred_domain = "devops" if "devops" in file.name.lower() else (
                              "datascience" if "ds" in file.name.lower() else (
                              "java" if "java" in file.name.lower() else (
                              "fullstack" if "full" in file.name.lower() else "general")))

            try:
                analysis, meta, resume_id = analyze_resume(text, inferred_domain)
                doc = Document(
                    page_content=text,
                    metadata={
                        "resume_id": resume_id,
                        "file_name": file.name,
                        "domain": inferred_domain,
                        **meta
                    }
                )
                resume_docs.append(doc)

                all_results.append({
                    "Resume ID": resume_id,
                    "File Name": file.name,
                    "Name": meta["name"],
                    "Phone": meta["phone"],
                    "Email": meta["email"],
                    "LinkedIn": meta["linkedin"],
                    "GitHub": meta["github"],
                    "Domain": inferred_domain,
                    "Summary": analysis.summary,
                    "Strengths": "\n".join(analysis.strengths),
                    "Weaknesses": "\n".join(analysis.weaknesses),
                    "Score": analysis.score,
                })

            except Exception as e:
                st.error(f"Error analyzing {file.name}: {e}")

        if job_desc and resume_docs:
            st.info("\U0001F50D Matching resumes to job description...")

            domain_guess = "devops" if "devops" in job_desc.lower() else (
                           "datascience" if "data" in job_desc.lower() else (
                           "java" if "java" in job_desc.lower() else (
                           "fullstack" if "full" in job_desc.lower() else "general")))

            vectorstore = build_vectorstore(resume_docs)
            top_docs = match_resumes(vectorstore, job_desc, domain_guess, k=TOP_K)
            top_ids = [doc.metadata["resume_id"] for doc in top_docs]

            evaluated = {}
            for doc in top_docs:
                try:
                    evaluation = evaluate_resume_vs_jd(doc.page_content, job_desc)
                    evaluated[doc.metadata["resume_id"]] = evaluation
                except Exception as e:
                    st.warning(f"Evaluation failed for {doc.metadata['file_name']}: {e}")

            for row in all_results:
                rid = row["Resume ID"]
                row["Job Match"] = "✅ Yes" if rid in top_ids else "❌ No"
                if rid in evaluated:
                    row["Strengths"] = "\n".join(evaluated[rid].strengths)
                    row["Weaknesses"] = "\n".join(evaluated[rid].weaknesses)
                    row["Score"] = evaluated[rid].score

        df = pd.DataFrame(all_results)
        st.success("✅ Analysis Complete")
        st.dataframe(df, use_container_width=True)

        excel_buffer = BytesIO()
        df.to_excel(excel_buffer, index=False, engine='openpyxl')
        excel_buffer.seek(0)

        st.download_button(
            label="📥 Download as Excel",
            data=excel_buffer,
            file_name="resume_analysis.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
